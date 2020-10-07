#!/usr/bin/env python
# encoding: utf-8
import json
import os
from flask import Flask, request, jsonify, render_template, redirect, url_for, send_from_directory
from os.path import join, dirname, realpath
from docx import Document

app = Flask(__name__, template_folder='template')


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def remove_row(table, row):
    tbl = table._tbl
    tr = row._tr
    tbl.remove(tr)


@app.errorhandler(404)
def page_not_found(error):
    return 'This page does not exist', 404


@app.route('/', methods=['GET', 'POST'])
def home():
    if request.method == 'POST' and request.form['submit_button'] == 'Download':
        resume = os.path.join(app.root_path, 'resume')
        file_path = os.path.join(resume, "Resume2.docx")
        doc_obj = Document(file_path)
        for paragraph in doc_obj.paragraphs:
            delete_paragraph(paragraph)
        for table in doc_obj.tables:
            for row in table.rows:
                remove_row(table, row)
        doc_obj.add_heading("Resume", 0)
        doc_obj.add_paragraph(request.form['name'])
        doc_obj.add_paragraph(request.form['email'])
        doc_obj.add_heading("Skills", 1)
        doc_obj.add_paragraph(request.form['skills'])
        doc_obj.add_heading("Projects", 1)
        doc_obj.add_paragraph(request.form['projects'])
        doc_obj.add_heading("Education", 1)
        doc_obj.add_paragraph(request.form['edu'])
        doc_obj.save(file_path)
        old_records = []
        new_record = request.form
        record_not_found = False
        if os.stat("data.txt").st_size != 0:
            with open('data.txt', 'r') as f:
                data = f.read()
            old_records = json.loads(data[:])
            for record in old_records:
                if record['name'] == new_record['name'] or record['email'] == new_record['email']:
                    record_not_found = False
                else:
                    record_not_found = True
            if record_not_found:
                old_records.append(new_record)
                with open('data.txt', 'w') as f:
                    f.write(json.dumps(old_records, indent=2))
            else:
                print("Info': 'name or email already exists")
        else:
            old_records.append(new_record)
            with open('data.txt', 'w') as f:
                f.write(json.dumps(old_records, indent=2))
        return send_from_directory(directory=resume, filename='Resume2.docx')
    return render_template("home.html")


@app.route("/about")
def about():
    return render_template("about.html")


if __name__ == "__main__":
    app.run(debug=True, port=8080)
