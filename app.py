#!/usr/bin/env python
# encoding: utf-8
import json
import os
from flask import Flask, request, jsonify, render_template, redirect, url_for, send_from_directory,flash
from os.path import join, dirname, realpath
from docx import Document
from forms import DownloadForm

# set a 'SECRET_KEY' to enable the Flask session cookies
app.config['SECRET_KEY'] = '12345687643'
app = Flask(__name__, template_folder='template')


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def remove_row(table, row):
    tbl = table._tbl
    tr = row._tr
    tbl.remove(tr)


@app.errorhandler(404)
def page_not_found(error):
    return 'This page does not exist', 404


@app.route('/', methods=('GET', 'POST'))
def submit():
    form = DownloadForm()
    print(form.errors)
    if request.method == 'GET':
        return render_template('submit.html', form=form)
    elif request.method == 'POST' and form.validate_on_submit():
        # form = DownloadForm(request.form)
        resume = os.path.join(app.root_path, 'resume')
        file_path = os.path.join(resume, "Resume2.docx")
        doc_obj = Document(file_path)
        for paragraph in doc_obj.paragraphs:
            delete_paragraph(paragraph)
        for table in doc_obj.tables:
            for row in table.rows:
                remove_row(table, row)
        doc_obj.add_heading("Resume", 0)
        doc_obj.add_paragraph(form.name.data)
        doc_obj.add_paragraph(form.email.data)
        doc_obj.add_heading("Skills", 1)
        doc_obj.add_paragraph(form.skills.data)
        doc_obj.add_heading("Projects", 1)
        doc_obj.add_paragraph(form.projects.data)
        doc_obj.add_heading("Education", 1)
        doc_obj.add_paragraph(form.education.data)
        doc_obj.save(file_path)
        old_records = []
        # new_record = form
        record_not_found = False
        if os.stat("data.txt").st_size != 0:
            with open('data.txt', 'r') as f:
                data = f.read()
            old_records = json.loads(data[:])
            for record in old_records:
                if record['name'] == form.name.data or record['email'] == form.email.data:
                    record_not_found = False
                else:
                    record_not_found = True
            if record_not_found:
                old_records.append(form.data)
                with open('data.txt', 'w') as f:
                    f.write(json.dumps(old_records, indent=2))
            else:
                print("Info': 'name or email already exists")
        else:
            old_records.append(form.data)
            with open('data.txt', 'w') as f:
                f.write(json.dumps(old_records, indent=2))
        return send_from_directory(directory=resume, filename='Resume2.docx')
    else:
        flash("Form is not valid")
        return render_template('submit.html', form=form)


@app.route("/about")
def about():
    return render_template("about.html")


if __name__ == "__main__":
    SECRET_KEY = os.urandom(32)
    app.config['SECRET_KEY'] = SECRET_KEY
    app.run(debug=True, port=8080)
